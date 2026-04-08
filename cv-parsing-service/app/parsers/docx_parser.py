"""
DOCX Parser - Extract text from Word documents
"""
from docx import Document
import logging

logger = logging.getLogger(__name__)


class DOCXParser:
    """Parse DOCX files and extract text"""
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            result = '\n'.join(paragraphs)
            logger.info(f"Successfully extracted from DOCX ({len(paragraphs)} paragraphs)")
            return result
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise
